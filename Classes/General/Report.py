from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io 

class Report():
    def __init__(self,Visual,Path,Text):
        self.Visual=Visual
        
        self.document = Document()
        self.document.add_heading('Stocks comparative periodical analysis',0)
        self.export_path=Path.tree.path["Visual"]+Path.slash+'report.docx'
        Text.append(self.export_path,"export_path")
        self.document.add_paragraph(Text.wrap())#toc
        self.add_figures()    
    def add_figures(self):
        self.add_fig("plot_bars()")
        self.add_fig("plot_box()")
        self.add_fig("plot_timeline('unNormalized','DeltaPerc')")
        self.add_fig("plot_timeline('Normalized','Investment Close')")
    def add_fig(self,plot_name):
        buf = io.BytesIO()
        plt.figure()
        eval("self.Visual."+plot_name)
        plt.savefig(buf, format='png')
        self.document.add_picture(buf, width=Inches(5.25))
        buf.close()
    def save(self):
        self.document.save(self.export_path)